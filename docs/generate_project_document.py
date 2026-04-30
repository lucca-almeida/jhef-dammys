from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(161, 77, 45)
TEXT = RGBColor(51, 51, 51)
MUTED = RGBColor(110, 110, 110)
LIGHT = "F6EFEA"
HEADER = "D9B7A6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="C8B7AD", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def style_run(run, size=11, bold=False, color=TEXT, font_name="Aptos"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_paragraph(doc, text="", size=11, bold=False, color=TEXT, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        style_run(run, size=11)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    style_run(run, size=16 if level == 1 else 13, bold=True, color=ACCENT)
    return paragraph


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        if widths:
            cell.width = Cm(widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, HEADER)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(header)
        style_run(run, size=10.5, bold=True, color=RGBColor(58, 33, 24))

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            if widths:
                cell.width = Cm(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows) % 2 == 0:
                set_cell_shading(cell, LIGHT)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(str(value))
            style_run(run, size=10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_cover(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    add_paragraph(doc, "DOCUMENTO TECNICO INICIAL", size=22, bold=True, color=ACCENT, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "Sistema interno para gestao de eventos, orcamentos, custos, estoque e lucro", size=13, color=MUTED, space_after=20, align=WD_ALIGN_PARAGRAPH.CENTER)

    box = doc.add_table(rows=4, cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.style = "Table Grid"
    set_table_borders(box, color="BFA89B")
    labels = ["Projeto", "Cliente interno", "Versao", "Data"]
    values = [
        "Sistema de administracao para servicos de buffet e churrasco",
        "Operacao propria do negocio",
        "V1 - planejamento inicial",
        "29/04/2026",
    ]
    for i, (label, value) in enumerate(zip(labels, values)):
        left = box.rows[i].cells[0]
        right = box.rows[i].cells[1]
        set_cell_shading(left, HEADER)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = left.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(label)
        style_run(r1, size=10.5, bold=True, color=RGBColor(58, 33, 24))
        p2 = right.paragraphs[0]
        r2 = p2.add_run(value)
        style_run(r2, size=10.5)

    add_paragraph(doc, "", space_after=18)
    add_paragraph(
        doc,
        "Este documento consolida a visao do sistema, o escopo do MVP, os modulos principais, a arquitetura sugerida e a ordem recomendada de desenvolvimento.",
        size=11.5,
        color=TEXT,
        space_after=8,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Objetivo central: transformar um processo hoje espalhado entre Instagram, WhatsApp, papel e memoria em uma operacao organizada, mensuravel e preparada para crescer.",
        size=11.5,
        color=TEXT,
        space_after=8,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    doc.add_page_break()


def build_document():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    styles["Normal"].font.size = Pt(11)

    add_cover(doc)

    add_heading(doc, "1. Visao geral do projeto")
    add_paragraph(
        doc,
        "O sistema sera criado inicialmente para uso interno, com foco em dar controle sobre agenda, orcamentos, eventos, custos, estoque, compras e lucro real.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "A fase 1 nao sera voltada ao cliente final. O foco e organizar a operacao diaria do negocio, reduzir retrabalho e substituir anotacoes no papel por um painel administrativo confiavel.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "2. Problemas atuais")
    add_bullets(
        doc,
        [
            "Atendimento espalhado entre Instagram e WhatsApp.",
            "Orcamentos montados no improviso, sem historico centralizado.",
            "Agenda suscetivel a confusao e perda de datas.",
            "Custos calculados no olho, sem visao clara do lucro real.",
            "Compras e estoque sem controle estruturado.",
            "Variacao de preco de mercado sem registro historico facil de consultar.",
        ],
    )

    add_heading(doc, "3. Objetivos da versao 1")
    add_bullets(
        doc,
        [
            "Organizar os pedidos desde o primeiro contato.",
            "Controlar agenda e eventos fechados em um unico lugar.",
            "Padronizar a criacao de orcamentos.",
            "Registrar pagamentos, custos e lucro por evento.",
            "Criar uma base de estoque e compras que permita melhor tomada de decisao.",
        ],
    )

    add_heading(doc, "4. Escopo do MVP")
    add_table(
        doc,
        ["Categoria", "Itens incluidos na V1"],
        [
            ["Essencial", "Login, clientes, servicos, orcamentos, eventos, agenda, pagamentos, custos e lucro"],
            ["Importante", "Dashboard, estoque basico, compras e historico de precos"],
            ["Futuro", "Area do cliente, consulta publica de datas, contrato automatico e automacoes"],
        ],
        widths=[4.2, 11.5],
    )

    add_heading(doc, "5. Requisitos funcionais principais")
    add_bullets(
        doc,
        [
            "Cadastrar clientes com nome, telefone, Instagram e observacoes.",
            "Cadastrar servicos e itens do cardapio com valor base.",
            "Criar orcamentos com quantidade de pessoas, data, local e itens escolhidos.",
            "Definir se o atendimento e so mao de obra ou material mais mao de obra.",
            "Converter orcamento aprovado em evento sem redigitar informacoes.",
            "Exibir os eventos em calendario para evitar conflitos de agenda.",
            "Registrar sinal, pagamentos parciais e pagamento final.",
            "Lancar custos por evento e calcular lucro real.",
            "Controlar estoque basico e registrar compras com variacao de preco.",
        ],
    )

    add_heading(doc, "6. Requisitos nao funcionais")
    add_bullets(
        doc,
        [
            "Interface simples, clara e facil de usar no dia a dia.",
            "Arquitetura separada entre frontend, backend e banco de dados.",
            "Seguranca com login, protecao de rotas e senhas criptografadas.",
            "Estrutura preparada para crescimento sem precisar refazer a base.",
            "Bom desempenho nas rotinas mais frequentes: agenda, cadastro, consulta e fechamento.",
            "Responsividade para uso em computador e consulta eventual no celular.",
        ],
    )

    add_heading(doc, "7. Fluxo principal do sistema")
    add_table(
        doc,
        ["Etapa", "Descricao"],
        [
            ["1", "Cliente entra em contato pelo Instagram ou WhatsApp"],
            ["2", "Cliente e cadastrado no sistema"],
            ["3", "Orcamento e criado com dados do evento e itens desejados"],
            ["4", "Orcamento recebe status e pode ser aprovado ou recusado"],
            ["5", "Orcamento aprovado vira evento na agenda"],
            ["6", "Pagamentos e custos sao registrados durante a execucao"],
            ["7", "Sistema calcula o lucro real ao final do evento"],
        ],
        widths=[2.0, 13.7],
    )

    add_heading(doc, "8. Entidades principais")
    add_table(
        doc,
        ["Entidade", "Papel no sistema"],
        [
            ["User", "Usuarios com acesso ao sistema"],
            ["Client", "Clientes que pedem orcamento ou fecham eventos"],
            ["Service", "Servicos e itens do cardapio"],
            ["Budget", "Orcamentos em negociacao"],
            ["BudgetItem", "Itens detalhados de cada orcamento"],
            ["Event", "Eventos confirmados e presentes na agenda"],
            ["Payment", "Sinal, pagamentos parciais e quitaçao"],
            ["EventCost", "Custos ligados a um evento especifico"],
            ["StockItem", "Itens de estoque e insumos"],
            ["Purchase", "Compras realizadas e historico de preco"],
        ],
        widths=[4.0, 11.7],
    )

    add_heading(doc, "9. Modulos da aplicacao")
    add_table(
        doc,
        ["Modulo", "Responsabilidade principal"],
        [
            ["Autenticacao", "Login, sessao e protecao das rotas"],
            ["Clientes", "Cadastro, busca e historico de atendimento"],
            ["Servicos", "Cardapio, itens e valores base"],
            ["Orcamentos", "Criacao, edicao e aprovacao de propostas"],
            ["Eventos", "Agenda real, dados do servico e status do evento"],
            ["Pagamentos", "Controle do sinal e saldo a receber"],
            ["Custos", "Registro de gastos por evento"],
            ["Dashboard", "Resumo operacional e financeiro"],
            ["Estoque", "Saldo de itens, entradas e saidas"],
            ["Compras", "Historico de aquisicoes e variacao de preco"],
            ["Relatorios", "Faturamento, gastos e lucro por periodo"],
        ],
        widths=[4.0, 11.7],
    )

    add_heading(doc, "10. Arquitetura recomendada")
    add_paragraph(
        doc,
        "A arquitetura sugerida separa a aplicacao em frontend, backend e banco de dados. Isso melhora a manutencao, facilita a evolucao para novas fases e reduz o risco de acoplamento excessivo.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_table(
        doc,
        ["Camada", "Tecnologia sugerida", "Motivo"],
        [
            ["Frontend", "Next.js + TypeScript + Tailwind", "Interface moderna, organizada e preparada para crescer"],
            ["Backend", "NestJS + TypeScript", "Estrutura em modulos, validacao e regras de negocio bem separadas"],
            ["Banco", "PostgreSQL", "Confiavel, robusto e adequado para sistema profissional"],
            ["ORM", "Prisma", "Modelagem clara, migrations e integracao forte com TypeScript"],
        ],
        widths=[3.0, 6.1, 6.4],
    )

    add_heading(doc, "11. Estrutura sugerida de pastas")
    add_paragraph(doc, "Organizacao macro do projeto:", bold=True, space_after=4)
    add_paragraph(
        doc,
        "new-project/\n  frontend/\n  backend/",
        size=10.5,
        color=RGBColor(70, 70, 70),
        space_after=8,
    )
    add_paragraph(doc, "Principais modulos previstos no backend:", bold=True, space_after=4)
    add_paragraph(
        doc,
        "auth, users, clients, services, budgets, events, payments, costs, stock, purchases, dashboard e reports.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    add_heading(doc, "12. Ordem recomendada de desenvolvimento")
    add_table(
        doc,
        ["Etapa", "Entrega"],
        [
            ["1", "Preparacao do projeto: frontend, backend, dependencias e estrutura de pastas"],
            ["2", "Modelagem do banco e criacao das entidades principais"],
            ["3", "Backend base com Prisma, validacao e configuracoes comuns"],
            ["4", "Autenticacao e usuario administrador"],
            ["5", "Modulo de clientes"],
            ["6", "Modulo de servicos"],
            ["7", "Modulo de orcamentos"],
            ["8", "Modulo de eventos e agenda"],
            ["9", "Pagamentos"],
            ["10", "Custos e calculo de lucro"],
            ["11", "Dashboard"],
            ["12", "Estoque, compras e relatorios"],
        ],
        widths=[2.0, 13.7],
    )

    add_heading(doc, "13. Roadmap futuro")
    add_bullets(
        doc,
        [
            "Area do cliente para consulta de disponibilidade e orcamento.",
            "Geracao automatica de contrato.",
            "Aceite digital.",
            "Automacoes para resposta e acompanhamento.",
            "Possivel controle ampliado de equipe e ajudantes.",
        ],
    )

    add_heading(doc, "14. Conclusao")
    add_paragraph(
        doc,
        "A versao 1 deve priorizar organizacao operacional. Antes de pensar em portal do cliente, automacoes ou contrato automatico, o sistema precisa resolver as dores internas: agenda, orcamentos, custos, estoque e lucro real.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        doc,
        "Com essa base bem feita, a fase 2 podera ser construida com muito mais seguranca e aproveitando a mesma arquitetura.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    section = doc.sections[-1]
    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Cm(16.0))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = footer_table.rows[0].cells
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(left.paragraphs[0].add_run("Sistema interno de eventos e buffet"), size=9, color=MUTED)
    style_run(right.paragraphs[0].add_run("Documento tecnico inicial"), size=9, color=MUTED)

    return doc


if __name__ == "__main__":
    output_dir = Path("docs")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "documento-tecnico-inicial.docx"
    build_document().save(output_path)
    print(output_path.resolve())
